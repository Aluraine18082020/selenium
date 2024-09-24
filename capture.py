import os
from docx import Document
from docx.shared import Inches

# Directory where Allure screenshots are stored
allure_results_dir = 'allure-results'
screenshots_dir = allure_results_dir  # Replace with actual screenshot path if different

# Create a new Word document
doc = Document()

# Title of the document
doc.add_heading('Allure Test Report Screenshots', 0)

# Loop through the screenshots in the Allure results folder
for filename in os.listdir(screenshots_dir):
    if filename.endswith('.png'):  # Look for PNG screenshots
        # Add a new section for each screenshot
        doc.add_heading(f"Screenshot: {filename}", level=1)

        # Add the screenshot image to the document
        screenshot_path = os.path.join(screenshots_dir, filename)
        doc.add_picture(screenshot_path, width=Inches(5.0))  # Adjust the width as needed

# Save the document
doc.save('allure_screenshots_report.docx')

print("Word document created successfully with screenshots.")
